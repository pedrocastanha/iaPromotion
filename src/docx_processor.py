from typing import List
from docx import Document as DocxDocument

class DocxProcessor:
    def __init__(self):
        pass

    def extract_text_from_docx(self, docx_path: str) -> str:
        try:
            doc = DocxDocument(docx_path)
            parts: List[str] = []

            for p in doc.paragraphs:
                txt = (p.text or "").strip()
                if txt:
                    parts.append(txt)

            for table in doc.tables:
                for row in table.rows:
                    cells = [(c.text or "").strip() for c in row.cells]
                    row_text = " | ".join([c for c in cells if c])
                    if row_text:
                        parts.append(row_text)
            return "\n".join(parts).strip()
        except Exception as e:
            raise Exception(f"Error extracting text from file {docx_path}: {e}")